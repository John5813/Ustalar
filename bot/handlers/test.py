import io
import logging

from aiogram import F, Router
from aiogram.fsm.context import FSMContext
from aiogram.types import BufferedInputFile, CallbackQuery, Message

from bot.keyboards import get_test_confirm_keyboard, get_test_format_keyboard, get_test_question_count_keyboard
from bot.states import TestStates
from database.database import Database
from services.ai_service import generate_test_questions
from translations import get_text

logger = logging.getLogger(__name__)
router = Router()

TEST_PRICE_PER_QUESTION = 500  # so'm per question


def _test_price(count: int) -> int:
    return count * TEST_PRICE_PER_QUESTION


def _build_test_docx(topic: str, questions: list, language: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    labels = {
        "uz": ("Test", "Javoblar", "Izoh"),
        "ru": ("Тест", "Ответы", "Пояснение"),
        "en": ("Test", "Answers", "Explanation"),
    }
    title_word, answers_word, explanation_word = labels.get(language, labels["uz"])

    title = doc.add_heading(f"{title_word}: {topic}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    letters = ["A", "B", "C", "D"]

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {q.get('question', '')}")
        run.bold = True
        run.font.size = Pt(12)
        for j, opt in enumerate(q.get("options", [])[:4]):
            doc.add_paragraph(f"   {letters[j]}) {opt}", style="List Bullet")
        doc.add_paragraph()

    doc.add_page_break()
    ans_heading = doc.add_heading(answers_word, level=2)
    ans_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for i, q in enumerate(questions, 1):
        idx = q.get("correct_index", 0)
        options = q.get("options", [])
        letter = letters[idx] if idx < 4 else str(idx)
        text = options[idx] if idx < len(options) else ""
        explanation = q.get("explanation", "")

        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {letter}) {text}")
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        if explanation:
            p2 = doc.add_paragraph()
            p2.add_run(f"   {explanation_word}: ").italic = True
            p2.add_run(explanation)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Entry ────────────────────────────────────────────────────────────────────

@router.callback_query(F.data == "os:test")
async def test_start(call: CallbackQuery, state: FSMContext, user_lang: str = "uz"):
    await state.set_state(TestStates.waiting_for_topic)
    await call.message.edit_text(
        get_text(user_lang, "test_enter_topic"),
        parse_mode="HTML",
    )
    await call.answer()


@router.callback_query(F.data == "os:test_back")
async def test_back_to_topic(call: CallbackQuery, state: FSMContext, user_lang: str = "uz"):
    await state.set_state(TestStates.waiting_for_topic)
    await call.message.edit_text(
        get_text(user_lang, "test_enter_topic"),
        parse_mode="HTML",
    )
    await call.answer()


# ─── Topic ────────────────────────────────────────────────────────────────────

@router.message(TestStates.waiting_for_topic)
async def test_got_topic(message: Message, state: FSMContext, user_lang: str = "uz"):
    topic = message.text.strip() if message.text else ""
    if len(topic) < 3:
        await message.answer(get_text(user_lang, "topic_too_short"))
        return
    await state.update_data(test_topic=topic)
    await state.set_state(TestStates.waiting_for_question_count)
    await message.answer(
        get_text(user_lang, "test_select_count"),
        reply_markup=get_test_question_count_keyboard(user_lang),
    )


# ─── Question count ───────────────────────────────────────────────────────────

@router.callback_query(F.data.startswith("test_count_"), TestStates.waiting_for_question_count)
async def test_got_count(call: CallbackQuery, state: FSMContext, user_lang: str = "uz"):
    count = int(call.data.split("_")[-1])
    price = _test_price(count)
    await state.update_data(test_count=count, test_price=price)
    await state.set_state(TestStates.waiting_for_format)

    price_labels = {
        "uz": f"💰 Narx: <b>{price:,} so'm</b> ({count} ta savol)\n\n",
        "ru": f"💰 Цена: <b>{price:,} сум</b> ({count} вопросов)\n\n",
        "en": f"💰 Price: <b>{price:,} sum</b> ({count} questions)\n\n",
    }
    prefix = price_labels.get(user_lang, price_labels["uz"])

    await call.message.edit_text(
        prefix + get_text(user_lang, "test_select_format"),
        reply_markup=get_test_format_keyboard(user_lang),
        parse_mode="HTML",
    )
    await call.answer()


@router.callback_query(F.data == "test_format_back")
async def test_format_back(call: CallbackQuery, state: FSMContext, user_lang: str = "uz"):
    await state.set_state(TestStates.waiting_for_question_count)
    await call.message.edit_text(
        get_text(user_lang, "test_select_count"),
        reply_markup=get_test_question_count_keyboard(user_lang),
    )
    await call.answer()


# ─── Format selected → confirm ────────────────────────────────────────────────

@router.callback_query(F.data.in_({"test_format_file", "test_format_poll"}), TestStates.waiting_for_format)
async def test_got_format(call: CallbackQuery, state: FSMContext, db: Database, user_lang: str = "uz"):
    data = await state.get_data()
    fmt = "file" if call.data == "test_format_file" else "poll"
    count = data.get("test_count", 10)
    price = data.get("test_price", _test_price(count))
    topic = data.get("test_topic", "")

    await state.update_data(test_format=fmt)

    user = await db.get_user(call.from_user.id)
    balance = user.balance if user else 0

    if balance < price:
        insuf = {
            "uz": f"❌ Hisobingizda mablag' yetarli emas.\n\n💰 Balans: <b>{balance:,} so'm</b>\n💳 Kerak: <b>{price:,} so'm</b>",
            "ru": f"❌ Недостаточно средств.\n\n💰 Баланс: <b>{balance:,} сум</b>\n💳 Нужно: <b>{price:,} сум</b>",
            "en": f"❌ Insufficient balance.\n\n💰 Balance: <b>{balance:,} sum</b>\n💳 Required: <b>{price:,} sum</b>",
        }.get(user_lang, "❌ Insufficient balance.")
        await call.message.edit_text(insuf, parse_mode="HTML")
        await call.answer()
        return

    fmt_uz = "Fayl (DOCX)" if fmt == "file" else "So'rovnoma (Poll)"
    fmt_ru = "Файл (DOCX)" if fmt == "file" else "Опрос (Poll)"
    fmt_en = "File (DOCX)" if fmt == "file" else "Poll"
    confirm_text = {
        "uz": (f"📝 <b>{topic}</b>\n\n"
               f"🔢 Savollar: <b>{count} ta</b>\n"
               f"📋 Format: <b>{fmt_uz}</b>\n"
               f"💰 Narx: <b>{price:,} so'm</b>\n\n"
               f"✅ Tasdiqlaysizmi?"),
        "ru": (f"📝 <b>{topic}</b>\n\n"
               f"🔢 Вопросов: <b>{count}</b>\n"
               f"📋 Формат: <b>{fmt_ru}</b>\n"
               f"💰 Цена: <b>{price:,} сум</b>\n\n"
               f"Подтверждаете?"),
        "en": (f"📝 <b>{topic}</b>\n\n"
               f"🔢 Questions: <b>{count}</b>\n"
               f"📋 Format: <b>{fmt_en}</b>\n"
               f"💰 Price: <b>{price:,} sum</b>\n\n"
               f"Confirm?"),
    }.get(user_lang, "")

    await call.message.edit_text(
        confirm_text,
        reply_markup=get_test_confirm_keyboard(user_lang),
        parse_mode="HTML",
    )
    await call.answer()


# ─── Cancel ───────────────────────────────────────────────────────────────────

@router.callback_query(F.data == "test_cancel")
async def test_cancel(call: CallbackQuery, state: FSMContext, user_lang: str = "uz"):
    await state.clear()
    cancel_text = {"uz": "❌ Bekor qilindi.", "ru": "❌ Отменено.", "en": "❌ Cancelled."}.get(user_lang, "❌")
    await call.message.edit_text(cancel_text)
    await call.answer()


# ─── Confirm → generate ───────────────────────────────────────────────────────

@router.callback_query(F.data == "test_confirm")
async def test_confirm_handler(call: CallbackQuery, state: FSMContext, db: Database, user_lang: str = "uz"):
    data = await state.get_data()
    topic = data.get("test_topic", "")
    count = data.get("test_count", 10)
    price = data.get("test_price", _test_price(count))
    fmt = data.get("test_format", "file")

    user = await db.get_user(call.from_user.id)
    if not user or user.balance < price:
        await call.answer("❌ Yetarli mablag' yo'q.", show_alert=True)
        return

    await db.update_user_balance(call.from_user.id, -price)
    await call.message.edit_text(get_text(user_lang, "test_generating"), parse_mode="HTML")
    await call.answer()
    await state.set_state(TestStates.generating)

    questions = await generate_test_questions(topic, count, user_lang)

    if not questions:
        await db.update_user_balance(call.from_user.id, price)
        await call.message.edit_text(get_text(user_lang, "test_error"))
        await state.clear()
        return

    try:
        if fmt == "file":
            docx_bytes = _build_test_docx(topic, questions, user_lang)
            safe_name = topic[:30].replace(" ", "_").replace("/", "_")
            file = BufferedInputFile(docx_bytes, filename=f"test_{safe_name}.docx")
            caption = get_text(user_lang, "test_ready_file").format(topic=topic, count=len(questions))
            await call.message.answer_document(file, caption=caption, parse_mode="HTML")
            await call.message.delete()
        else:
            letters = ["A", "B", "C", "D"]
            for i, q in enumerate(questions, 1):
                question_text = f"❓ {i}/{len(questions)}: {q.get('question', '')}"
                options = q.get("options", [])[:4]
                correct_idx = q.get("correct_index", 0)
                explanation = q.get("explanation", "")

                await call.bot.send_poll(
                    chat_id=call.from_user.id,
                    question=question_text[:300],
                    options=[f"{letters[j]}) {opt}"[:100] for j, opt in enumerate(options)],
                    type="quiz",
                    correct_option_id=correct_idx,
                    explanation=explanation[:200] if explanation else None,
                    is_anonymous=True,
                )

            done_text = get_text(user_lang, "test_ready_poll").format(count=len(questions), topic=topic)
            await call.message.edit_text(done_text, parse_mode="HTML")

        logger.info(f"Test generated: user={call.from_user.id}, topic={topic}, count={count}, fmt={fmt}")
    except Exception as e:
        logger.error(f"Test delivery error: {e}")
        await db.update_user_balance(call.from_user.id, price)
        await call.message.edit_text(get_text(user_lang, "test_error"))

    await state.clear()
