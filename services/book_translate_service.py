import logging
import os
import re
import asyncio
from openai import AsyncOpenAI
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import TEMP_DIR

logger = logging.getLogger(__name__)

GEMINI_MODEL = "google/gemini-2.0-flash-001"
CHUNK_WORD_LIMIT = 4000
SECTION_WORD_LIMIT = 750

LANG_NAMES = {
    "uz": "Uzbek",
    "ru": "Russian",
    "en": "English",
}

LANG_SUFFIXES = {
    "uz": "_uz",
    "ru": "_ru",
    "en": "_en",
}


def _get_client() -> AsyncOpenAI:
    return AsyncOpenAI(
        api_key=os.environ.get("AI_INTEGRATIONS_OPENROUTER_API_KEY") or "dummy-key",
        base_url=os.environ.get("AI_INTEGRATIONS_OPENROUTER_BASE_URL"),
    )


def _word_count(text: str) -> int:
    return len(text.split())


def count_docx_words(file_path: str) -> int:
    doc = Document(file_path)
    total = 0
    for para in doc.paragraphs:
        if para.text.strip():
            total += _word_count(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        total += _word_count(para.text)
    return total


def get_book_translate_price(word_count: int) -> int:
    from config import BOOK_TRANSLATE_PRICES
    for limit, price in sorted(BOOK_TRANSLATE_PRICES.items()):
        if word_count <= limit:
            return price
    return list(BOOK_TRANSLATE_PRICES.values())[-1]


WORDS_PER_PAGE = 300  # Bir standart kitob varogi uchun so'zlar soni

# Russian-specific Cyrillic letters that don't appear in Uzbek
_RU_SPECIFIC = re.compile(r'[ыЫёЁэЭъЪ]')
# General Cyrillic
_CYRILLIC = re.compile(r'[а-яА-ЯёЁ]')
# Latin
_LATIN = re.compile(r'[a-zA-Z]')


def detect_source_language(file_path: str) -> str:
    """Hujjatning asosiy tilini aniqlaydi: 'ru', 'uz' yoki 'en'.
    
    Birinchi 80 ta paragrafni namunaviy o'qiydi:
    - Ko'p rus-xos belgilar (ы, ё, э, ъ) → 'ru'
    - Ko'p kirill → 'ru' (lekin o'zbek kirilli ham bo'lishi mumkin)
    - Ko'p lotin → 'en' yoki 'uz' (latin)
    """
    try:
        doc = Document(file_path)
        sample = " ".join(p.text for p in doc.paragraphs[:80] if p.text.strip())[:5000]
        ru_specific = len(_RU_SPECIFIC.findall(sample))
        cyrillic = len(_CYRILLIC.findall(sample))
        latin = len(_LATIN.findall(sample))
        # Agar rus-xos harflar ko'p bo'lsa → albatta rus
        if ru_specific > 10:
            return "ru"
        # Kirill ustun bo'lsa → rus yoki o'zbek kirill
        if cyrillic > latin * 1.5:
            return "ru"
        # Lotin ustun → ingliz yoki o'zbek lotin
        return "en"
    except Exception:
        return "unknown"


def count_estimated_pages(file_path: str) -> int:
    """Kitobdagi taxminiy varoq sonini hisoblaydi (300 so'z = 1 varoq)."""
    words = count_docx_words(file_path)
    return max(1, round(words / WORDS_PER_PAGE))


def get_page_range_word_count(file_path: str, from_page: int, to_page: int) -> int:
    """Berilgan varoq oralig'idagi so'zlar sonini hisoblaydi."""
    doc = Document(file_path)
    start_word = (from_page - 1) * WORDS_PER_PAGE
    end_word = to_page * WORDS_PER_PAGE

    current_word = 0
    selected_words = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        wlen = _word_count(text)
        if current_word + wlen > start_word and current_word < end_word:
            overlap = min(current_word + wlen, end_word) - max(current_word, start_word)
            selected_words += max(0, overlap)
        current_word += wlen
        if current_word >= end_word:
            break
    return selected_words


def extract_pages_by_range(input_path: str, from_page: int, to_page: int) -> str:
    """Berilgan varoq oralig'iga mos paragraflarni yangi DOCX ga ajratib oladi."""
    from copy import deepcopy
    doc = Document(input_path)
    start_word = (from_page - 1) * WORDS_PER_PAGE
    end_word = to_page * WORDS_PER_PAGE

    current_word = 0
    selected_paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        wlen = _word_count(text) if text else 0
        if current_word + wlen > start_word and current_word < end_word:
            selected_paras.append(para)
        current_word += wlen
        if current_word >= end_word:
            break

    new_doc = Document()
    for p in list(new_doc.paragraphs):
        p._element.getparent().remove(p._element)
    for para in selected_paras:
        new_doc.element.body.append(deepcopy(para._element))

    base = os.path.splitext(os.path.basename(input_path))[0]
    out_name = f"bt_pages_{from_page}_{to_page}_{base}.docx"
    out_path = os.path.join(TEMP_DIR, out_name)
    new_doc.save(out_path)
    return out_path


async def extract_toc_from_docx(file_path: str) -> list[dict]:
    """Extract table of contents from a DOCX file using AI.

    Returns list of dicts: [{"number": 1, "title": "Chapter title"}, ...]
    """
    doc = Document(file_path)
    all_text = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            all_text.append(text)
        if len(all_text) > 600:
            break

    sample = "\n".join(all_text[:500])

    client = _get_client()
    prompt = (
        "You are an expert at analyzing book structures. "
        "Below is the text from a book document. "
        "Find and list ALL chapters, sections, or major headings from the table of contents or document structure.\n\n"
        "Rules:\n"
        "- Return ONLY a numbered list of chapter/section titles\n"
        "- Include ALL chapters/sections you can find\n"
        "- Format: one per line, like: 1. Chapter Title\n"
        "- If there's a clear table of contents, use it\n"
        "- If no clear TOC, identify major headings/sections from the text structure\n"
        "- Do NOT add explanations\n\n"
        f"DOCUMENT TEXT:\n{sample}"
    )

    response = await client.chat.completions.create(
        model=GEMINI_MODEL,
        messages=[
            {"role": "system", "content": "You analyze book structures. Return only numbered chapter lists."},
            {"role": "user", "content": prompt},
        ],
        max_tokens=2000,
        temperature=0.2,
    )
    raw = response.choices[0].message.content.strip()

    chapters = []
    for line in raw.split("\n"):
        line = line.strip()
        if not line:
            continue
        match = re.match(r'^(\d+)\s*[.\):\-]\s*(.+)', line)
        if match:
            num = int(match.group(1))
            title = match.group(2).strip()
            chapters.append({"number": num, "title": title})

    if not chapters:
        chapters = [{"number": 1, "title": line.strip()} for i, line in enumerate(raw.split("\n"), 1) if line.strip()]
        for i, ch in enumerate(chapters):
            ch["number"] = i + 1

    return chapters


def _normalize(text: str) -> str:
    """Normalize text for matching: lowercase, collapse whitespace, remove punctuation."""
    import unicodedata
    text = text.lower().strip()
    text = re.sub(r'[^\w\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _similarity_score(title: str, paragraph: str) -> float:
    """Score how well a chapter title matches a paragraph (0.0 to 1.0)."""
    t_norm = _normalize(title)
    p_norm = _normalize(paragraph)

    if not t_norm or not p_norm:
        return 0.0

    if t_norm in p_norm:
        return 1.0

    t_words = t_norm.split()
    p_words = set(p_norm.split())

    if not t_words:
        return 0.0

    matched = sum(1 for w in t_words if w in p_words)
    ratio = matched / len(t_words)

    if len(paragraph.strip()) < 200 and ratio >= 0.6:
        return ratio
    elif ratio >= 0.8:
        return ratio

    return 0.0


def _find_chapter_boundaries(doc: Document, toc: list[dict], selected_numbers: list[int] = None) -> dict[int, tuple[int, int]]:
    """Find paragraph index boundaries for each chapter in the document.

    Returns: {chapter_number: (start_para_idx, end_para_idx)}
    Uses fuzzy title matching, earliest match per chapter, and monotonically increasing boundaries.
    """
    chapter_titles = {ch["number"]: ch["title"] for ch in toc}
    para_count = len(doc.paragraphs)

    first_match: dict[int, tuple[int, float]] = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) > 500:
            continue
        for num, title in chapter_titles.items():
            if num in first_match:
                continue
            score = _similarity_score(title, text)
            if score >= 0.6:
                first_match[num] = (i, score)
                logger.debug(f"Chapter {num} matched at para {i} (score={score:.2f}): '{text[:80]}'")

    if not first_match:
        logger.warning(f"No chapter boundaries found via title matching. "
                       f"Tried {len(chapter_titles)} titles against {para_count} paragraphs.")

    chapter_starts = sorted(first_match.items(), key=lambda x: x[1][0])

    prev_pos = -1
    filtered = []
    for num, (pos, score) in chapter_starts:
        if pos > prev_pos:
            filtered.append((num, pos))
            prev_pos = pos

    boundaries = {}
    for idx, (num, start) in enumerate(filtered):
        if selected_numbers and num not in selected_numbers:
            continue
        if idx + 1 < len(filtered):
            end = filtered[idx + 1][1]
        else:
            end = para_count
        boundaries[num] = (start, end)

    return boundaries


def count_chapters_words(file_path: str, toc: list[dict], selected_numbers: list[int]) -> int:
    """Count words only in the selected chapters."""
    doc = Document(file_path)
    boundaries = _find_chapter_boundaries(doc, toc, selected_numbers)

    total = 0
    for num in selected_numbers:
        if num in boundaries:
            start, end = boundaries[num]
            for i in range(start, min(end, len(doc.paragraphs))):
                text = doc.paragraphs[i].text.strip()
                if text:
                    total += _word_count(text)

    if total == 0:
        total_all = count_docx_words(file_path)
        ratio = len(selected_numbers) / max(len(toc), 1)
        total = int(total_all * ratio)

    return total


def _split_into_chunks(texts: list[str], limit: int = CHUNK_WORD_LIMIT) -> list[list[str]]:
    chunks: list[list[str]] = []
    current: list[str] = []
    current_words = 0
    for t in texts:
        wc = _word_count(t)
        if current_words + wc > limit and current:
            chunks.append(current)
            current = [t]
            current_words = wc
        else:
            current.append(t)
            current_words += wc
    if current:
        chunks.append(current)
    return chunks


class TranslationError(Exception):
    pass


async def _extract_glossary(client: AsyncOpenAI, doc, target_lang: str) -> str:
    """Extract a terminology glossary from the first ~5000 words of the document.

    Returns a compact glossary string (e.g. 'John Smith → Jon Smit') to be
    included in every chunk's translation prompt for consistent naming.
    Returns an empty string on failure or when no glossary is needed.
    """
    lang_name = LANG_NAMES.get(target_lang, "English")

    _GLOSSARY_WORD_TARGET = 5000
    _GLOSSARY_CHAR_CAP = 32000

    words_collected = 0
    text_parts: list[str] = []

    def _collect(text: str) -> bool:
        nonlocal words_collected
        text = text.strip()
        if text:
            text_parts.append(text)
            words_collected += len(text.split())
        return words_collected < _GLOSSARY_WORD_TARGET

    for para in doc.paragraphs:
        if not _collect(para.text):
            break

    if words_collected < _GLOSSARY_WORD_TARGET:
        for table in doc.tables:
            if words_collected >= _GLOSSARY_WORD_TARGET:
                break
            for row in table.rows:
                if words_collected >= _GLOSSARY_WORD_TARGET:
                    break
                for cell in row.cells:
                    if words_collected >= _GLOSSARY_WORD_TARGET:
                        break
                    for para in cell.paragraphs:
                        if not _collect(para.text):
                            break

    sample_text = "\n".join(text_parts)
    if not sample_text or words_collected < 200:
        return ""

    prompt = (
        f"Read the following text excerpt and create a short terminology glossary for translating it into {lang_name}. "
        f"Identify: character/person names, place names, organization names, and domain-specific terms that appear more than once. "
        f"For each, provide its {lang_name} translation or transliteration. "
        f"Format: one entry per line — 'Original → Translation'. "
        f"Maximum 40 entries. Output ONLY the glossary lines, nothing else. "
        f"If the source text is already in {lang_name}, respond with exactly: N/A\n\n"
        f"TEXT:\n{sample_text[:_GLOSSARY_CHAR_CAP]}"
    )

    try:
        response = await client.chat.completions.create(
            model=GEMINI_MODEL,
            messages=[
                {"role": "system", "content": f"You are a skilled {lang_name} author preparing to rewrite a text. Extract key names and terms that must stay consistent. Be concise and accurate."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=800,
            temperature=0.1,
        )
        result = response.choices[0].message.content.strip()
        if not result or result == "N/A" or len(result) < 10:
            return ""
        logger.info(f"Glossary extracted: {len(result)} chars, ~{result.count(chr(10)) + 1} entries")
        return result
    except Exception as e:
        logger.warning(f"Glossary extraction failed (non-fatal): {e}")
        return ""


async def _rewrite_section(
    client: AsyncOpenAI,
    section_texts: list[str],
    target_lang: str,
    prev_context: str = "",
    glossary: str = "",
    max_retries: int = 2,
) -> str:
    """Read a 2-3 page section as a whole and rewrite it as coherent academic prose.

    Instead of paragraph-by-paragraph numbered mapping, the entire section is
    sent as one block. The AI understands it fully and writes flowing, logically
    connected output. Returns a string with paragraphs separated by double newlines.
    """
    lang_name = LANG_NAMES.get(target_lang, "English")
    block = "\n\n".join(section_texts)

    if target_lang == "uz":
        script_note = (
            " IMPORTANT: Write EXCLUSIVELY in Uzbek Latin script. "
            "Do NOT use Cyrillic letters at all. Use proper Uzbek Latin alphabet: "
            "a, b, d, e, f, g, h, i, j, k, l, m, n, o, p, q, r, s, t, u, v, x, y, z, "
            "o', g', sh, ch, ng. Example: 'bo'ladi', 'o'quvchi', 'g'oya'."
        )
    else:
        script_note = ""

    glossary_note = (
        f"\n\nNAME & TERM GLOSSARY (use these consistently throughout):\n{glossary}"
        if glossary else ""
    )

    prev_note = (
        f"\n\nPREVIOUS SECTION END (already written — do NOT repeat; "
        f"your text must continue naturally from this point):\n{prev_context}"
        if prev_context else ""
    )

    system_prompt = (
        f"You are a skilled {lang_name} academic writer.{script_note}\n"
        f"You will receive a 2-3 page excerpt of source text (possibly in another language or poorly written). "
        f"Your task:\n"
        f"1. Read and fully understand the entire excerpt as a unit.\n"
        f"2. Rewrite it as clear, fluent, academic {lang_name} prose — NOT a word-for-word translation.\n"
        f"3. Express the full meaning naturally, as a native {lang_name} author communicating these ideas to a reader.\n"
        f"4. Write in continuous paragraphs separated by a blank line. Each paragraph must logically "
        f"flow into the next — the whole output should read as one coherent piece of text.\n"
        f"5. Skip any corrupted or unreadable fragments (OCR artifacts, broken symbols) — do not reproduce them.\n"
        f"Output ONLY the rewritten {lang_name} text. No labels, no headings, no commentary."
        f"{glossary_note}"
        f"{prev_note}"
    )

    last_error = None
    for attempt in range(max_retries + 1):
        try:
            response = await client.chat.completions.create(
                model=GEMINI_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": block},
                ],
                max_tokens=16000,
                temperature=0.5,
            )
            result = response.choices[0].message.content.strip()
            return result
        except Exception as e:
            last_error = e
            logger.warning(f"Section rewrite attempt {attempt + 1} failed: {e}")
            if attempt < max_retries:
                await asyncio.sleep(2 ** attempt)

    raise TranslationError(f"Section rewrite failed after {max_retries + 1} attempts: {last_error}")


async def _translate_chunk(
    client: AsyncOpenAI,
    texts: list[str],
    target_lang: str,
    prev_context: str = "",
    glossary: str = "",
    max_retries: int = 2,
) -> list[str]:
    lang_name = LANG_NAMES.get(target_lang, "English")
    numbered = "\n".join(f"[{i}] {t}" for i, t in enumerate(texts))

    if target_lang == "uz":
        script_note = (
            " IMPORTANT: Write EXCLUSIVELY in Uzbek Latin script. "
            "Do NOT use Cyrillic letters at all. Use proper Uzbek Latin alphabet: "
            "a, b, d, e, f, g, h, i, j, k, l, m, n, o, p, q, r, s, t, u, v, x, y, z, "
            "o', g', sh, ch, ng. Example: 'bo'ladi', 'o'quvchi', 'g'oya'."
        )
    else:
        script_note = ""

    glossary_note = (
        f"\n\nNAME & TERM GLOSSARY (use these consistently in your rewriting):\n{glossary}"
        if glossary else ""
    )

    context_note = (
        f"\n\nPREVIOUS SECTION (already rewritten — do NOT rewrite again; "
        f"use it only to maintain continuity of narrative, tone, and any unfinished thoughts):\n{prev_context}"
        if prev_context else ""
    )

    system_prompt = (
        f"You are a skilled {lang_name} writer and author.{script_note} "
        f"Your task: read and deeply understand each numbered text segment, then rewrite it naturally in {lang_name} "
        f"as if you are a native {lang_name} author who originally wrote this content — NOT a translator. "
        f"Express the same ideas, meaning, and intent, but use natural {lang_name} phrasing, sentence flow, and style. "
        "Do NOT translate word-for-word. Write the way a fluent native speaker would express these ideas. "
        "Return ONLY the rewritten segments, each prefixed with its original number in square brackets, "
        "like [0], [1], etc. Do not add any explanations or commentary. "
        "If a segment contains corrupted or unreadable characters (OCR scan artifacts like random symbols, "
        "broken words with '^', '*', digits mixed with letters), rewrite only the clearly readable parts "
        "and skip the unreadable fragments — do not reproduce the corruption."
        f"{glossary_note}"
        f"{context_note}"
    )

    last_error = None
    for attempt in range(max_retries + 1):
        try:
            response = await client.chat.completions.create(
                model=GEMINI_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": numbered},
                ],
                max_tokens=16000,
                temperature=0.5,
            )
            raw = response.choices[0].message.content.strip()
            return _parse_numbered_response(raw, len(texts))
        except Exception as e:
            last_error = e
            logger.warning(f"Translation chunk attempt {attempt + 1} failed: {e}")
            if attempt < max_retries:
                await asyncio.sleep(2 ** attempt)

    raise TranslationError(f"Translation failed after {max_retries + 1} attempts: {last_error}")


def _parse_numbered_response(raw: str, expected: int) -> list[str]:
    parts = re.split(r'\[(\d+)\]', raw)
    mapping: dict[int, str] = {}
    i = 1
    while i < len(parts) - 1:
        idx_str = parts[i]
        text = parts[i + 1].strip()
        try:
            mapping[int(idx_str)] = text
        except ValueError:
            pass
        i += 2
    result = []
    for j in range(expected):
        result.append(mapping.get(j, ""))
    return result


def _clean_paragraph_text(text: str) -> str:
    """Strip footnote markers and problematic characters from paragraph text."""
    # Unicode superscript digits: ¹²³⁴⁵⁶⁷⁸⁹⁰ (also subscript range)
    text = re.sub(r'[\u00b9\u00b2\u00b3\u2070-\u2079\u2080-\u2089]+', '', text)
    # [1], [2] or (1), (2) style footnote markers
    text = re.sub(r'\[\d+\]|\(\d+\)', '', text)
    # Non-printable control characters (keep \t and \n)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # Collapse multiple spaces
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()


def _strip_footnotes_from_doc(doc) -> None:
    """Remove all footnote/endnote reference markers from every paragraph in doc."""
    _FN_TAGS = {qn('w:footnoteReference'), qn('w:endnoteReference')}

    def _strip_para(para):
        p = para._p
        for elem in list(p.iter()):
            if elem.tag in _FN_TAGS:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)

    for para in doc.paragraphs:
        _strip_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _strip_para(para)


def _looks_like_heading(text: str) -> bool:
    """Heuristic: short line with no sentence-ending punctuation → heading."""
    words = text.split()
    if len(words) > 9:
        return False
    if text and text[-1] in '.?!;,':
        return False
    return True


def _set_para_xml_format(p, is_heading: bool) -> None:
    """Apply paragraph-level alignment and first-line indent via raw XML.

    Headings  → centered, no indent.
    Body text → both-justified, first-line indent 720 twips (~1.27 cm).
    """
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)

    # Alignment
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center' if is_heading else 'both')

    # First-line indent (body only; remove for headings)
    ind = pPr.find(qn('w:ind'))
    if is_heading:
        if ind is not None:
            pPr.remove(ind)
    else:
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLine'), '720')


def _apply_para_format(para, is_heading: bool) -> None:
    """Format a python-docx paragraph object (alignment, indent, font).

    Used for newly-created paragraphs (add_paragraph) where python-docx
    Run objects are accessible directly.
    Headings  → centered, 14 pt bold Times New Roman, no indent.
    Body text → justified, 12 pt Times New Roman, first-line indent 1.25 cm.
    """
    pf = para.paragraph_format
    if is_heading:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = None
        target_size = Pt(14)
        bold = True
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.25)
        target_size = Pt(12)
        bold = False

    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = target_size
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)


def _replace_paragraph_text(para, new_text: str):
    """Replace paragraph text with a single clean Times New Roman 12pt run.

    Instead of distributing text across original runs (which inherits
    superscript/color/tiny-font formatting and produces garbled output),
    this clears ALL run elements and writes the translated text as one
    clean run. Paragraph-level style (Heading/Normal) is kept intact.
    """
    p = para._p

    # Run tags that signal an inline image/object — these runs must stay in place
    _IMAGE_CONTENT_TAGS = {qn('w:drawing'), qn('w:object'), qn('w:pict')}

    # Non-run tags to always remove
    _REMOVE_OTHER = {
        qn('w:hyperlink'), qn('w:del'), qn('w:ins'),
        qn('w:footnoteReference'), qn('w:endnoteReference'),
        qn('w:proofErr'), qn('w:bookmarkStart'), qn('w:bookmarkEnd'),
    }

    # Identify the first text-only run (marks insertion position for new run)
    # and collect all other non-image elements for removal.
    first_text_run = None
    to_remove = []
    for child in list(p):
        if child.tag == qn('w:r'):
            has_image = any(elem.tag in _IMAGE_CONTENT_TAGS for elem in child.iter())
            if has_image:
                pass  # preserve image runs in their original position
            elif first_text_run is None:
                first_text_run = child  # placeholder — replaced below
            else:
                to_remove.append(child)
        elif child.tag in _REMOVE_OTHER:
            to_remove.append(child)

    for elem in to_remove:
        p.remove(elem)

    # Detect heading paragraphs (bold + centered + 14pt)
    style_name = ''
    try:
        style_name = (para.style.name or '').lower()
    except Exception:
        pass
    is_heading = 'heading' in style_name

    # Apply paragraph-level alignment and first-line indent
    _set_para_xml_format(p, is_heading)

    # Build one clean run: Times New Roman, black, no superscript
    # Headings → 14pt bold; body → 12pt normal
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.append(rFonts)

    # heading=14pt (28 half-points), body=12pt (24 half-points)
    pt_val = '28' if is_heading else '24'
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), pt_val)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), pt_val)
    rPr.append(sz)
    rPr.append(szCs)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)

    if is_heading:
        rPr.append(OxmlElement('w:b'))

    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = new_text
    if new_text != new_text.strip():
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)

    # Insert the new run at the position of the first original text run so that
    # any inline image runs that follow it remain in their original positions.
    # (addprevious inserts r as the preceding sibling of first_text_run; then we
    # remove the old placeholder run.)
    if first_text_run is not None:
        first_text_run.addprevious(r)
        p.remove(first_text_run)
    else:
        p.append(r)


def _extract_chapters_to_new_doc(input_path: str, toc: list[dict], selected_chapters: list[int]) -> str | None:
    """Extract only the selected chapters into a brand-new DOCX file.

    Returns the path to the new DOCX, or None if no chapter boundaries were found.
    """
    from copy import deepcopy

    source_doc = Document(input_path)
    boundaries = _find_chapter_boundaries(source_doc, toc, selected_chapters)

    if not boundaries:
        logger.warning(f"No boundaries found for chapters {selected_chapters}. Falling back to full doc.")
        return None

    allowed_para_indices = set()
    for num in selected_chapters:
        if num in boundaries:
            start, end = boundaries[num]
            for i in range(start, min(end, len(source_doc.paragraphs))):
                allowed_para_indices.add(i)

    if not allowed_para_indices:
        logger.warning("No paragraphs found in chapter boundaries.")
        return None

    logger.info(f"Extracting {len(allowed_para_indices)} paragraphs for chapters {selected_chapters}")

    new_doc = Document()
    body = new_doc.element.body
    body.clear()

    for i in sorted(allowed_para_indices):
        para = source_doc.paragraphs[i]
        new_elem = deepcopy(para._element)
        body.append(new_elem)

    base, ext = os.path.splitext(os.path.basename(input_path))
    ch_label = "_".join(str(n) for n in sorted(selected_chapters))
    out_name = f"{base}_ch{ch_label}{ext}"
    out_path = os.path.join(TEMP_DIR, out_name)
    os.makedirs(TEMP_DIR, exist_ok=True)
    new_doc.save(out_path)
    return out_path


async def translate_docx(input_path: str, target_language: str, selected_chapters: list[int] = None, toc: list[dict] = None) -> str:
    """Rewrite a DOCX file in the target language using section-based academic rewriting.

    The document body is split into ~2-3 page sections (SECTION_WORD_LIMIT words each).
    Each section is sent to the AI as a whole block: the AI reads and understands it
    fully, then rewrites it as coherent, fluent academic prose in the target language.
    Sections are chained with prev_context so the output reads as one continuous text.
    A new DOCX is built from the rewritten paragraphs.

    If selected_chapters is given, extracts those chapters into a separate DOCX first.
    """
    client = _get_client()

    chapters_extracted = False
    if selected_chapters and toc:
        extracted_path = _extract_chapters_to_new_doc(input_path, toc, selected_chapters)
        if extracted_path:
            working_path = extracted_path
            doc = Document(working_path)
            chapters_extracted = True
            logger.info(f"Extracted chapters {selected_chapters} to {working_path}")
        else:
            logger.warning("Chapter extraction failed — processing full document instead")
            working_path = input_path
            doc = Document(input_path)
            selected_chapters = None
    else:
        working_path = input_path
        doc = Document(input_path)

    _strip_footnotes_from_doc(doc)

    para_texts: list[str] = []
    for para in doc.paragraphs:
        text = _clean_paragraph_text(para.text.strip())
        if text:
            para_texts.append(text)

    glossary = await _extract_glossary(client, doc, target_language)

    rewritten_blocks: list[str] = []
    prev_context = ""
    if para_texts:
        sections = _split_into_chunks(para_texts, limit=SECTION_WORD_LIMIT)
        for sec_idx, section in enumerate(sections):
            wc = sum(_word_count(t) for t in section)
            logger.info(f"Rewriting section {sec_idx + 1}/{len(sections)} ({wc} words, {len(section)} paragraphs)")
            rewritten = await _rewrite_section(
                client, section, target_language,
                prev_context=prev_context,
                glossary=glossary,
            )
            rewritten_blocks.append(rewritten)
            output_paras = [p.strip() for p in rewritten.split("\n\n") if p.strip()]
            prev_context = output_paras[-1] if output_paras else ""

    out_doc = Document()
    if out_doc.paragraphs:
        default_p = out_doc.paragraphs[0]._element
        default_p.getparent().remove(default_p)

    for block in rewritten_blocks:
        output_paras = [p.strip() for p in block.split("\n\n") if p.strip()]
        for para_text in output_paras:
            is_h = _looks_like_heading(para_text)
            p = out_doc.add_paragraph(para_text, style="Normal")
            _apply_para_format(p, is_h)

    cell_refs: list[tuple] = []
    cell_texts: list[str] = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = _clean_paragraph_text(para.text.strip())
                    if text:
                        cell_refs.append((t_idx, r_idx, c_idx, p_idx))
                        cell_texts.append(text)

    if cell_texts:
        from copy import deepcopy
        for table in doc.tables:
            out_doc.element.body.append(deepcopy(table._tbl))

        cell_chunks = _split_into_chunks(cell_texts)
        translated_cells: list[str] = []
        cell_prev_ctx = ""
        for chunk_idx, chunk in enumerate(cell_chunks):
            logger.info(f"Rewriting table chunk {chunk_idx + 1}/{len(cell_chunks)}")
            translated = await _translate_chunk(
                client, chunk, target_language,
                prev_context=cell_prev_ctx,
                glossary=glossary,
            )
            translated_cells.extend(translated)
            ctx_segs = translated[-4:] if len(translated) >= 4 else translated
            cell_prev_ctx = "\n".join(s for s in ctx_segs if s.strip())

        for i, (t_idx, r_idx, c_idx, p_idx) in enumerate(cell_refs):
            new_text = _clean_paragraph_text(translated_cells[i] if i < len(translated_cells) else "")
            if not new_text:
                continue
            para = out_doc.tables[t_idx].rows[r_idx].cells[c_idx].paragraphs[p_idx]
            _replace_paragraph_text(para, new_text)

    base, ext = os.path.splitext(os.path.basename(input_path))
    suffix = LANG_SUFFIXES.get(target_language, f"_{target_language}")
    if selected_chapters:
        ch_label = "_".join(str(n) for n in sorted(selected_chapters))
        out_name = f"{base}_ch{ch_label}{suffix}{ext}"
    else:
        out_name = f"{base}{suffix}{ext}"
    out_path = os.path.join(TEMP_DIR, out_name)
    os.makedirs(TEMP_DIR, exist_ok=True)
    out_doc.save(out_path)

    if chapters_extracted and working_path != input_path:
        try:
            os.remove(working_path)
        except Exception:
            pass

    return out_path


async def auto_convert_pdf_to_docx(pdf_path: str) -> str:
    """Convert PDF to DOCX using converter_service, return DOCX path."""
    from services.converter_service import convert_pdf_to_docx
    return await convert_pdf_to_docx(pdf_path)


async def extract_relevant_content_for_topic(book_content: str, topic: str, lang: str = "uz") -> str:
    """
    Kitob mazmunidan faqat berilgan mavzuga oid qismlarni ajratib oladi.
    Bu funksiya AI yordamida kitobni 'mavzu bo'yicha qidiradi' va eng aloqador
    bo'limlarni qaytaradi. Agar hech narsa topilmasa, asl mazmun qaytariladi.
    """
    if not book_content or len(book_content.strip()) < 100:
        return book_content

    lang_prompts = {
        "uz": (
            f"Quyidagi kitob matnidan '{topic}' mavzusiga bevosita aloqador bo'lgan "
            f"barcha paragraflar, bo'limlar va gaplarni ajratib ber. "
            f"FAQAT shu mavzuga tegishli qismlarni ko'chir — boshqa narsa yozma. "
            f"Agar mavzu bilan bog'liq bo'lim yo'q bo'lsa, eng yaqin mavzudagi qismni ber.\n\n"
            f"KITOB MATNI:\n{book_content[:50000]}"
        ),
        "ru": (
            f"Из следующего текста книги извлеки все абзацы и разделы, "
            f"которые непосредственно связаны с темой '{topic}'. "
            f"Выпиши ТОЛЬКО релевантные части — ничего лишнего. "
            f"Если нет точного совпадения, дай ближайшие по теме фрагменты.\n\n"
            f"ТЕКСТ КНИГИ:\n{book_content[:50000]}"
        ),
        "en": (
            f"From the following book text, extract all paragraphs and sections "
            f"that are directly related to the topic '{topic}'. "
            f"Copy ONLY the relevant parts — nothing else. "
            f"If no exact match, provide the closest topic-related excerpts.\n\n"
            f"BOOK TEXT:\n{book_content[:50000]}"
        ),
    }

    prompt = lang_prompts.get(lang, lang_prompts["uz"])

    try:
        client = _get_client()
        response = await client.chat.completions.create(
            model=GEMINI_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an expert text extractor. Your job is to find and extract "
                        "only the passages from a book that are relevant to a given topic. "
                        "Do not summarize, do not add anything — only copy relevant original text."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            max_tokens=6000,
            temperature=0.1,
        )
        extracted = response.choices[0].message.content.strip()

        min_words = max(500, len(book_content.split()) // 5)
        extracted_words = len(extracted.split())

        if extracted and extracted_words >= min_words:
            logger.info(f"Extracted {extracted_words} words relevant to topic: {topic!r}")
            return extracted
        elif extracted and extracted_words >= 100:
            # Ajratilgan matn kam, lekin to'liq matn bilan birlashtir
            logger.info(f"Extracted only {extracted_words} words, merging with full content")
            return extracted + "\n\n---\n\n" + book_content
        else:
            logger.warning(f"AI extraction too short ({extracted_words} words) for topic {topic!r}, using full content")
            return book_content
    except Exception as e:
        logger.error(f"Failed to extract topic-relevant content: {e}")
        return book_content
